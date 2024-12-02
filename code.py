import pymupdf
import re
import time
from googletrans import Translator
from transformers import pipeline
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


start_time = time.time()


def main(input_pdf):
    en_text, tables, title = extract_data(input_pdf)
    cleared_en_text = remove_extra_text(en_text, tables)
    split_en_text = text_partition(cleared_en_text)
    ru_text = translation(split_en_text)
    summary = summarization(split_en_text)
    return docx_write(split_en_text, ru_text, summary, title)


def extract_data(input_pdf):
    with pymupdf.open(input_pdf) as doc:  # open document
        en_text = chr(12).join([page.get_text() for page in doc])
        tables = []
        for page in doc:
            tabs = pymupdf.Page.find_tables(page)
            if tabs.tables:
                tables.append(tabs[0].extract())
    metadata = doc.metadata
    if len(tables) == 0:
        tables = []
    return en_text, tables, metadata["title"]


def delete_tables(en_text, tables_list):
    try:
        for table in tables_list:
            table_text = []
            for subpart in table:
                for text in subpart:
                    if text:
                        table_text.append(text)
            table_text = "\n".join(table_text)
            en_text = re.sub(fr"{table_text.strip()}", "", en_text)
    except re.error:
        return en_text
    return en_text


def remove_extra_text(text, tables_list):
    """
    Removes extra text e.g Figures, references, etc using regex and slices[:].
    """

    # pattern remove dashes "-"
    text = re.sub(r"-\s+", "", text)
    # pattern_strip lines
    text = re.sub(r" \s+", "\n", text)
    # pattern_before_abstract
    text = text[text.find('Abstract'):]  # Works
    # pattern_references
    text = text[:text.find("REFERENCES")]  # Works
    # pattern_remove_tables
    text = delete_tables(text, tables_list)  # Works
    # pattern_article_number
    text = re.sub(r'\b.*DOI:.*\n', "", text)
    # pattern_ieee_number
    text = re.sub(r'\n+.*IEEE\n', '\n', text)
    # pattern_ieee_number
    text = re.sub(r'\n+IEEE.*\n', '\n', text)
    # pattern conference
    text = re.sub(r'\n+.*(Conference|Proseedings|International Conference).*\n', '\n', text)
    # pattern_licence
    text = re.sub(r'\n\d+\nAuthorized.*\n', '', text)  # Works
    # pattern_licence
    text = re.sub(r'\nAuthorized.*\n', '', text)  # Works
    # pattern_tables remove titles
    text = re.sub(r'\bTABLE .*\n.*\n', '', text)
    # pattern delete figure annotation
    text = re.sub(r"\b(?:Fig|Figure)\.?\s+\d{1,2}\.\s.*\n*(?=[A-Z])?", "\n", text)
    # pattern restrictions
    text = re.sub(r"Restrictions apply", "", text)
    return text


def split_long_part(long_part):
    """
    Splits parts longer than 3000 cr (summarizer limit) in the middle
    """
    parts = []
    dot_index = len(long_part) // 2
    cr = ""
    dots = re.findall(r"[a-z]+\.", long_part)
    split_dot = None
    for dot in dots:
        if split_dot is None:
            split_dot = long_part.find(dot)
        elif abs(split_dot - dot_index) > abs(long_part.find(dot) - dot_index):
            split_dot = long_part.find(dot)
    while cr != ".":
        cr = long_part[split_dot]
        split_dot += 1
    parts.append(long_part[:split_dot])
    parts.append(long_part[split_dot + 1:])
    return parts


def text_partition(en_text):
    """
    Splits text in smaller parts based on the start of the paragraph e.g I. Introduction etc.
    """
    en_text = en_text[en_text.find('Abstract'):]  # Works
    # pattern_references
    en_text = en_text[:en_text.find("REFERENCES")]  # Works
    # pattern_extra_spaces
    split = re.findall(r"\n\b(?:XII|XI|X|IX|VIII|VII|VI|V|IV|III|II|I)\.?\s[A-Z]*", en_text)
    split.insert(0, "Abstract")
    # split and reverse text
    split_text = []
    stop_index = -1
    for i in split[::-1]:
        start_index = en_text.find(i)
        split_text.append(en_text[start_index:stop_index])
        stop_index = start_index
    split_text = split_text[::-1]
    en_text = []
    for part in split_text:
        # check if part is longer than 3000 cr
        if len(part) > 3000:
            split_part = split_long_part(part)
            for parted in split_part:
                en_text.append(parted)
        else:
            en_text.append(part)
    en_text_no_spaces = []
    for part in en_text:
        part_no_spaces = re.sub(r'\s+', ' ', part)
        en_text_no_spaces.append(part_no_spaces)
    return en_text_no_spaces


def translation(en_text):
    """
    Translate text from English to Russian using google translator API
    """
    ru_text = []
    translator = Translator()
    for part in en_text:
        trans = translator.translate(part, src='en', dest='ru')
        ru_text.append(trans.text)
    return ru_text


def summarization(en_text):
    """
    Summarize en_text using HuggingFace pre-trained model "facebook/bart-large-cnn"
    """
    summarizer = pipeline("summarization", model="facebook/bart-large-cnn")
    summary = ""
    # en_text = en_text.split("\n\n")
    for part in en_text:
        part_summary = summarizer(part, do_sample=False)
        summary += str(part_summary)
    return summary


def docx_write(en_text, ru_text, summary, title):
    """
    Create docx document, apply formatting. Add table with 2 columns |en_text|ru_text|,
    en_text length and summary
    """
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    no_spaces_text = re.sub(r"\s*", "", "".join(en_text))
    text_length = str(len(no_spaces_text))
    doc.add_paragraph(f"This article contains {text_length} symbols.")
    heading = doc.add_paragraph(title)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table_rows = len(en_text)
    table_columns = 2
    table = doc.add_table(table_rows, table_columns, style=None)
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    # write en_text in the first column using loop
    en_raw_number = 0
    for en_part in en_text:
        if en_raw_number > len(en_text):
            break
        else:
            table.cell(en_raw_number, 0).text = en_part
            en_raw_number += 1
    # write ru_text in the second column using loop
    ru_raw_number = 0
    for ru_part in ru_text:
        if ru_raw_number > len(en_text):
            break
        else:
            table.cell(ru_raw_number, 1).text = ru_part
            ru_raw_number += 1
    summary = summary.split("]")
    doc.add_paragraph()
    summ = doc.add_paragraph("SUMMARY")
    summ.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for part in summary:
        doc.add_paragraph(part[19:-2])
    doc.save(f"Uploaded files/{title}.docx")
    return f"Uploaded files/{title}.docx"
